"""Document classifier by composition: EXTRACT → ANALYZE → DECIDE.

Generic, format-agnostic classification pipeline.

- EXTRACT: pull raw signals from the document (per-format plugins)
- ANALYZE: compute derived features from raw signals (format-agnostic)
- DECIDE: classify based on features (rule-based, extensible)

Usage:
    classifier = DocumentClassifier()
    result = classifier.classify("path/to/file.pdf")

Each stage can be called independently for debugging:
    signals = classifier.extract("path/to/file.pdf")
    features = classifier.analyze(signals)
    result = classifier.decide(features)
"""

from __future__ import annotations

import logging
import re
from collections.abc import Callable
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path

logger = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════════════════════
# Models
# ═══════════════════════════════════════════════════════════════════════════


class DocumentType(str, Enum):
    """Document categories that determine the processing pipeline."""

    TEXT_HEAVY = "text_heavy"
    TABLE_HEAVY = "table_heavy"
    MIXED = "mixed"
    SCANNED = "scanned"


@dataclass
class RawSignals:
    """Raw signals extracted from a document (format-specific).

    Populated by the EXTRACT stage. Each extractor fills what it can;
    unknown fields stay at their defaults.
    """

    file_path: str = ""
    file_extension: str = ""
    page_count: int = 0
    chars_per_page: list[float] = field(default_factory=list)
    lines_per_page: list[int] = field(default_factory=list)
    table_lines_per_page: list[int] = field(default_factory=list)
    list_lines_per_page: list[int] = field(default_factory=list)
    images_per_page: list[int] = field(default_factory=list)
    scanned_pages: list[int] = field(default_factory=list)


@dataclass
class Features:
    """Derived features computed from raw signals (format-agnostic).

    Populated by the ANALYZE stage.
    """

    file_extension: str = ""
    page_count: int = 0
    avg_chars_per_page: float = 0.0
    avg_lines_per_page: float = 0.0
    table_ratio: float = 0.0
    list_ratio: float = 0.0
    scanned_ratio: float = 0.0
    has_images: bool = False
    has_tables: bool = False
    has_lists: bool = False
    text_density: float = 0.0  # avg chars per line


@dataclass(frozen=True)
class ClassificationResult:
    """Output of the classifier."""

    doc_type: DocumentType
    confidence: float
    features: Features


# ═══════════════════════════════════════════════════════════════════════════
# STAGE 1: EXTRACT — pull raw signals (per-format plugins)
# ═══════════════════════════════════════════════════════════════════════════

_SAMPLE_PAGES = 5
_MIN_CHARS_SCANNED = 50

# Heuristics for table-like lines
_TABLE_LINE_RE_PIPE = re.compile(r".*\|.*\|.*\|")
_TABLE_LINE_RE_TAB = re.compile(r".*\t.*\t")
_TABLE_LINE_RE_SPACES = re.compile(r" {3,}")
_LIST_LINE_RE = re.compile(r"^\s*[-*+]\s+|^\s*\d+[.\)]\s+")


def _is_table_line(line: str) -> bool:
    """Heuristic: does this line look like a table row?"""
    s = line.strip()
    if not s:
        return False
    if _TABLE_LINE_RE_PIPE.match(s):
        return True
    if _TABLE_LINE_RE_TAB.match(s):
        return True
    if len(_TABLE_LINE_RE_SPACES.findall(s)) >= 2:
        return True
    return False


def _is_list_line(line: str) -> bool:
    """Heuristic: does this line look like a list item?"""
    return bool(_LIST_LINE_RE.match(line.strip()))


def extract_pdf(path: Path) -> RawSignals:
    """Extract signals from a PDF using pymupdf."""
    import pymupdf  # noqa: PLC0415

    doc = pymupdf.open(str(path))
    page_count = len(doc)
    sample = min(page_count, _SAMPLE_PAGES)

    signals = RawSignals(
        file_path=str(path),
        file_extension=path.suffix.lower(),
        page_count=page_count,
    )

    for i in range(sample):
        page = doc[i]
        text = page.get_text("text") or ""
        char_count = len(text.strip())
        signals.chars_per_page.append(char_count)

        lines = [ln for ln in text.split("\n") if ln.strip()]
        signals.lines_per_page.append(len(lines))
        signals.table_lines_per_page.append(sum(1 for ln in lines if _is_table_line(ln)))
        signals.list_lines_per_page.append(sum(1 for ln in lines if _is_list_line(ln)))

        image_list = page.get_images(full=True)
        signals.images_per_page.append(len(image_list))

        if char_count < _MIN_CHARS_SCANNED and image_list:
            signals.scanned_pages.append(i + 1)

    doc.close()
    return signals


def extract_text_file(path: Path) -> RawSignals:
    """Extract signals from a plain text or markdown file."""
    text = path.read_text(encoding="utf-8", errors="replace")
    lines = [ln for ln in text.split("\n") if ln.strip()]

    return RawSignals(
        file_path=str(path),
        file_extension=path.suffix.lower(),
        page_count=1,
        chars_per_page=[len(text.strip())],
        lines_per_page=[len(lines)],
        table_lines_per_page=[sum(1 for ln in lines if _is_table_line(ln))],
        list_lines_per_page=[sum(1 for ln in lines if _is_list_line(ln))],
        images_per_page=[0],
    )


def extract_docx(path: Path) -> RawSignals:
    """Extract signals from a DOCX file."""
    try:
        import docx  # noqa: PLC0415

        doc = docx.Document(str(path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        lines = [ln for ln in full_text.split("\n") if ln.strip()]
        table_count = len(doc.tables)

        return RawSignals(
            file_path=str(path),
            file_extension=".docx",
            page_count=max(1, len(full_text) // 3000),  # estimate
            chars_per_page=[len(full_text)],
            lines_per_page=[len(lines)],
            table_lines_per_page=[table_count * 5],  # rough estimate
            list_lines_per_page=[sum(1 for ln in lines if _is_list_line(ln))],
            images_per_page=[0],
        )
    except ImportError:
        return extract_text_file(path)


# Extractor registry: extension → extractor function
_EXTRACTORS: dict[str, Callable[..., RawSignals]] = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".txt": extract_text_file,
    ".md": extract_text_file,
    ".html": extract_text_file,
    ".htm": extract_text_file,
}


def extract(path: str | Path) -> RawSignals:
    """Extract signals from any supported document format."""
    path = Path(path)
    ext = path.suffix.lower()
    extractor = _EXTRACTORS.get(ext, extract_text_file)
    try:
        return extractor(path)
    except Exception as exc:  # noqa: BLE001
        logger.warning("Extraction failed for %s: %s, using text fallback", path.name, exc)
        return extract_text_file(path)


# ═══════════════════════════════════════════════════════════════════════════
# STAGE 2: ANALYZE — compute derived features (format-agnostic)
# ═══════════════════════════════════════════════════════════════════════════


def analyze(signals: RawSignals) -> Features:
    """Compute classification features from raw signals."""
    total_lines = sum(signals.lines_per_page) or 1
    total_table_lines = sum(signals.table_lines_per_page)
    total_list_lines = sum(signals.list_lines_per_page)
    total_chars = sum(signals.chars_per_page) or 1
    n_pages = max(len(signals.chars_per_page), 1)
    sample_pages = max(len(signals.chars_per_page), 1)

    return Features(
        file_extension=signals.file_extension,
        page_count=signals.page_count,
        avg_chars_per_page=total_chars / n_pages,
        avg_lines_per_page=total_lines / n_pages,
        table_ratio=total_table_lines / total_lines,
        list_ratio=total_list_lines / total_lines,
        scanned_ratio=len(signals.scanned_pages) / sample_pages,
        has_images=any(n > 0 for n in signals.images_per_page),
        has_tables=total_table_lines > 2,
        has_lists=total_list_lines > 2,
        text_density=total_chars / total_lines,
    )


# ═══════════════════════════════════════════════════════════════════════════
# STAGE 3: DECIDE — classify based on features (rule-based)
# ═══════════════════════════════════════════════════════════════════════════

# Thresholds (configurable)
TABLE_HEAVY_THRESHOLD = 0.25
MIXED_THRESHOLD = 0.10
SCANNED_THRESHOLD = 0.50


def decide(features: Features) -> ClassificationResult:
    """Classify a document based on its features."""
    # Scanned: most sampled pages have very little text but images
    if features.scanned_ratio > SCANNED_THRESHOLD:
        confidence = min(0.6 + features.scanned_ratio * 0.3, 0.95)
        return ClassificationResult(DocumentType.SCANNED, confidence, features)

    # Table-heavy
    if features.table_ratio > TABLE_HEAVY_THRESHOLD:
        confidence = min(0.5 + features.table_ratio, 0.95)
        if features.avg_chars_per_page > 500 and features.table_ratio < 0.6:
            return ClassificationResult(DocumentType.MIXED, confidence * 0.9, features)
        return ClassificationResult(DocumentType.TABLE_HEAVY, confidence, features)

    # Mixed: some tables
    if features.table_ratio > MIXED_THRESHOLD:
        return ClassificationResult(DocumentType.MIXED, 0.7, features)

    # Default: text-heavy
    confidence = min(0.7 + (1 - features.table_ratio) * 0.25, 0.95)
    return ClassificationResult(DocumentType.TEXT_HEAVY, confidence, features)


# ═══════════════════════════════════════════════════════════════════════════
# Composed classifier
# ═══════════════════════════════════════════════════════════════════════════


class DocumentClassifier:
    """Generic document classifier: EXTRACT → ANALYZE → DECIDE.

    Each stage can be called independently for debugging::

        classifier = DocumentClassifier()
        signals = classifier.extract("file.pdf")
        features = classifier.analyze(signals)
        result = classifier.decide(features)

    Or all at once::

        result = classifier.classify("file.pdf")
    """

    def classify(self, file_path: str | Path) -> ClassificationResult:
        """Run full pipeline: EXTRACT → ANALYZE → DECIDE."""
        path = Path(file_path)
        signals = self.extract(path)
        features = self.analyze(signals)
        result = self.decide(features)

        logger.info(
            "Classified %s: type=%s confidence=%.2f (pages=%d, table_ratio=%.2f, "
            "scanned_ratio=%.2f, has_tables=%s, has_images=%s)",
            path.name,
            result.doc_type.value,
            result.confidence,
            features.page_count,
            features.table_ratio,
            features.scanned_ratio,
            features.has_tables,
            features.has_images,
        )
        return result

    @staticmethod
    def extract(path: str | Path) -> RawSignals:
        """Stage 1: extract raw signals from document."""
        return extract(path)

    @staticmethod
    def analyze(signals: RawSignals) -> Features:
        """Stage 2: compute features from raw signals."""
        return analyze(signals)

    @staticmethod
    def decide(features: Features) -> ClassificationResult:
        """Stage 3: classify based on features."""
        return decide(features)
