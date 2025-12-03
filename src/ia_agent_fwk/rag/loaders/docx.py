"""DOCX file loader using python-docx."""

from __future__ import annotations

import logging
from pathlib import Path
from typing import ClassVar

from ia_agent_fwk.rag.exceptions import DocumentLoadError
from ia_agent_fwk.rag.loaders.base import FileLoader
from ia_agent_fwk.rag.models import Document

logger = logging.getLogger(__name__)


class DOCXLoader(FileLoader):
    """Load ``.docx`` files using ``python-docx``.

    Requires the optional ``python-docx`` package.
    """

    supported_extensions: ClassVar[list[str]] = [".docx"]

    async def load(self, file_path: str | Path) -> Document:
        """Extract text from a DOCX file and return a ``Document``."""
        try:
            import docx
        except ImportError as exc:
            msg = "python-docx is required for DOCX loading. Install with: pip install python-docx"
            raise DocumentLoadError(msg) from exc

        path = Path(file_path)
        if not path.exists():
            msg = f"File not found: {path}"
            raise DocumentLoadError(msg)

        try:
            doc = docx.Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            content = "\n\n".join(paragraphs)
        except DocumentLoadError:
            raise
        except Exception as exc:
            msg = f"Failed to load DOCX file '{path}': {exc}"
            raise DocumentLoadError(msg) from exc

        return Document(
            content=content,
            metadata={
                "filename": path.name,
                "file_size_bytes": path.stat().st_size,
                "paragraph_count": len(paragraphs),
            },
            source=str(path),
            doc_type="docx",
        )
